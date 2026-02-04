#!/usr/bin/env python3
"""
extract_tables.py - Extract tables from specified chapters in DOCX files

Usage:
    python extract_tables.py <docx_file> --chapter <chapter_number> --output <output_file>
    
Example:
    python extract_tables.py document.docx --chapter 10 --output tables.json
"""

import argparse
import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError:
    print("Error: python-docx package is required")
    print("Please run: pip install python-docx")
    sys.exit(1)


def find_chapter_start(doc: Document, chapter_number: int) -> int:
    """
    Find the starting paragraph index of specified chapter
    
    Args:
        doc: Word document object
        chapter_number: Chapter number
        
    Returns:
        Paragraph index, -1 if not found
    """
    patterns = [
        rf'^{chapter_number}[\.\s]',          # "10." or "10 "
        rf'^Chapter\s*{chapter_number}\b',    # "Chapter 10"
        rf'^Section\s*{chapter_number}\b',    # "Section 10"
    ]
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        for pattern in patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return i
    return -1


def find_chapter_end(doc: Document, start_index: int, chapter_number: int) -> int:
    """
    Find the end position of chapter (start of next chapter or end of document)
    
    Args:
        doc: Word document object
        start_index: Current chapter start index
        chapter_number: Current chapter number
        
    Returns:
        End paragraph index
    """
    next_chapter = chapter_number + 1
    patterns = [
        rf'^{next_chapter}[\.\s]',
        rf'^Chapter\s*{next_chapter}\b',
        rf'^Section\s*{next_chapter}\b',
    ]
    
    for i in range(start_index + 1, len(doc.paragraphs)):
        text = doc.paragraphs[i].text.strip()
        for pattern in patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return i
    
    return len(doc.paragraphs)


def extract_table_data(table: Table) -> dict:
    """
    Extract table data
    
    Args:
        table: Word table object
        
    Returns:
        Dictionary containing headers and rows
    """
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)
    
    if not rows:
        return {"headers": [], "rows": []}
    
    return {
        "headers": rows[0],
        "rows": rows[1:] if len(rows) > 1 else []
    }


def get_chapter_title(doc: Document, para_index: int) -> str:
    """Get chapter title"""
    if 0 <= para_index < len(doc.paragraphs):
        return doc.paragraphs[para_index].text.strip()
    return ""


def extract_tables_from_chapter(docx_path: str, chapter_number: int) -> dict:
    """
    Extract all tables from specified chapter
    
    Args:
        docx_path: DOCX file path
        chapter_number: Chapter number
        
    Returns:
        Dictionary containing all table data
    """
    doc = Document(docx_path)
    
    # Find chapter range
    start_idx = find_chapter_start(doc, chapter_number)
    if start_idx == -1:
        return {
            "source_file": Path(docx_path).name,
            "chapter": f"Chapter {chapter_number} (not found)",
            "error": f"Chapter {chapter_number} not found",
            "tables": []
        }
    
    end_idx = find_chapter_end(doc, start_idx, chapter_number)
    chapter_title = get_chapter_title(doc, start_idx)
    
    # Get document XML structure to locate table positions
    # python-docx tables are document-wide, need to determine position
    tables_data = []
    table_index = 0
    
    # Iterate through document body elements
    body = doc.element.body
    current_para_idx = 0
    
    for element in body:
        if element.tag.endswith('p'):  # Paragraph
            current_para_idx += 1
        elif element.tag.endswith('tbl'):  # Table
            if start_idx <= current_para_idx < end_idx:
                # This table is in target chapter
                if table_index < len(doc.tables):
                    table = doc.tables[table_index]
                    table_data = extract_table_data(table)
                    table_data["index"] = len(tables_data) + 1
                    tables_data.append(table_data)
            table_index += 1
    
    return {
        "source_file": Path(docx_path).name,
        "chapter": chapter_title,
        "chapter_number": chapter_number,
        "tables": tables_data
    }


def extract_all_tables(docx_path: str) -> dict:
    """
    Extract all tables from document
    
    Args:
        docx_path: DOCX file path
        
    Returns:
        Dictionary containing all table data
    """
    doc = Document(docx_path)
    
    tables_data = []
    for i, table in enumerate(doc.tables):
        table_data = extract_table_data(table)
        table_data["index"] = i + 1
        tables_data.append(table_data)
    
    return {
        "source_file": Path(docx_path).name,
        "chapter": "All Document",
        "tables": tables_data
    }


def main():
    parser = argparse.ArgumentParser(
        description="Extract tables from DOCX files",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
    # Extract tables from chapter 10
    python extract_tables.py document.docx --chapter 10 --output tables.json
    
    # Extract all tables
    python extract_tables.py document.docx --output tables.json
        """
    )
    
    parser.add_argument("docx_file", help="DOCX file path")
    parser.add_argument("--chapter", "-c", type=int, help="Chapter number (if not specified, extract all)")
    parser.add_argument("--output", "-o", help="Output JSON file path (if not specified, output to stdout)")
    
    args = parser.parse_args()
    
    # Check if file exists
    if not Path(args.docx_file).exists():
        print(f"Error: File not found {args.docx_file}", file=sys.stderr)
        sys.exit(1)
    
    # Extract tables
    if args.chapter:
        result = extract_tables_from_chapter(args.docx_file, args.chapter)
    else:
        result = extract_all_tables(args.docx_file)
    
    # Output result
    output_json = json.dumps(result, ensure_ascii=False, indent=2)
    
    if args.output:
        Path(args.output).write_text(output_json, encoding="utf-8")
        print(f"Output saved to {args.output}")
        print(f"Found {len(result['tables'])} table(s)")
    else:
        print(output_json)


if __name__ == "__main__":
    main()
